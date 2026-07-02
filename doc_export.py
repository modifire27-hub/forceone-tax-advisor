# -*- coding: utf-8 -*-
"""
doc_export.py
==============
세무 자문 결과(마크다운 텍스트)를 docx / pdf 파일로 변환하는 모듈입니다.

- md  : 별도 변환 없이 그대로 .md 파일로 저장 (tax_advisor_engine.py의 save_response와 동일)
- docx: python-docx 사용. 마크다운 헤더(###)와 굵게(**) 표시를 기본적인 Word 서식으로 변환.
        본문 한글이 네모(□)로 깨지지 않도록, 모든 run에 w:eastAsia 폰트를 명시적으로 지정함.
- pdf : fpdf2 사용. 한글 출력을 위해 이 폴더의 assets/fonts/NanumGothic-Regular.ttf를
        최우선으로 사용(로컬/웹 배포 환경 무관하게 항상 사용 가능). 이 파일이 없으면
        로컬 Windows 시스템 폰트(맑은 고딕)나 Linux 시스템의 나눔고딕을 대신 찾아봄.
        폰트를 전혀 찾지 못하면 PDF 생성을 건너뛰고 명확한 안내 메시지를 반환.

Windows 환경 기준으로 작성됨.
"""

import re
from pathlib import Path
from datetime import datetime


# ----------------------------------------------------------------------
# 공통: 마크다운 본문을 줄 단위로 단순 파싱
# ----------------------------------------------------------------------
def _parse_markdown_lines(md_text: str):
    """
    마크다운 텍스트를 (종류, 내용) 튜플 리스트로 변환.
    종류: 'h3', 'h2', 'h1', 'bullet', 'text', 'blank'
    굵게(**...**) 표시는 내용 안에 마커를 남겨 호출부에서 처리.
    """
    lines = []
    for raw_line in md_text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            lines.append(("blank", ""))
        elif line.startswith("### "):
            lines.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            lines.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            lines.append(("h1", line[2:].strip()))
        elif line.strip().startswith("- "):
            lines.append(("bullet", line.strip()[2:].strip()))
        else:
            lines.append(("text", line.strip()))
    return lines


def _split_bold(text: str):
    """**굵게** 표시를 (텍스트, 굵음여부) 조각 리스트로 분리"""
    parts = []
    pattern = re.compile(r"\*\*(.+?)\*\*")
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            parts.append((text[last_end:m.start()], False))
        parts.append((m.group(1), True))
        last_end = m.end()
    if last_end < len(text):
        parts.append((text[last_end:], False))
    if not parts:
        parts = [(text, False)]
    return parts


# ----------------------------------------------------------------------
# DOCX 변환
# ----------------------------------------------------------------------
def build_docx_bytes(title: str, question: str, response_md: str) -> bytes:
    """
    질의/회신을 Word 문서(.docx) 형식의 바이트로 변환 (디스크에 쓰지 않음).

    설계 의도 (2026-06-25 — 웹 배포 지원):
    - 기존에는 output_path(파일 경로)를 받아 디스크에 직접 .save()하는 방식이었음.
    - PC 로컬 환경에서는 괜찮지만, 웹 서버(Streamlit Cloud)에는 사용자가 지정한
      로컬 경로(예: D드라이브 경로)가 존재하지 않아 디스크 쓰기 자체가 실패함.
    - 그래서 디스크에 쓰는 대신, 메모리(BytesIO)에 담아 바이트로 반환하도록 변경.
      이렇게 하면 호출하는 쪽(streamlit_ui.py)에서 st.download_button으로 바로
      사용자 컴퓨터에 다운로드시킬 수 있고, 동시에 로컬 환경에서도 그 바이트를
      파일로 저장하면 똑같이 동작함 (한 번 만들고 양쪽에 다 쓸 수 있는 구조).

    Parameters
    ----------
    title : str
        문서 제목
    question : str
        사용자 질문 (또는 종합 문서의 경우 안내문)
    response_md : str
        마크다운 형식의 AI 회신 본문

    Returns
    -------
    bytes
        .docx 파일의 바이트 내용
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ------------------------------------------------------------------
    # 디자인 테마 (streamlit_ui.py 화면 디자인과 색감을 통일)
    # 포인트 컬러: 짙은 블루(#2563EB), 제목 텍스트: 네이비(#1E3A5F),
    # 보조 텍스트: 중간 회색(#6B7280)
    # ------------------------------------------------------------------
    PF_ACCENT = RGBColor(0x25, 0x63, 0xEB)
    PF_TEXT_STRONG = RGBColor(0x1E, 0x3A, 0x5F)
    PF_TEXT_MUTED = RGBColor(0x6B, 0x72, 0x80)
    PF_BORDER_GRAY = "D1D5DB"

    def _set_paragraph_border_bottom(paragraph, color_hex="2563EB", size=12):
        """문단 아래쪽에 단일 보더(가로줄)를 추가하는 저수준 OOXML 조작 헬퍼.
        python-docx는 문단 보더를 직접 지원하지 않아 XML을 수동으로 삽입함."""
        p_pr = paragraph._p.get_or_add_pPr()
        p_borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), color_hex)
        p_borders.append(bottom)
        p_pr.append(p_borders)

    def _apply_korean_font(run, name="맑은 고딕"):
        """
        run에 한글 폰트를 적용.

        원인 메모 (2026-07-02, DOCX 본문이 네모(□)로 깨지는 문제 조사):
        python-docx의 run.font.name = "맑은 고딕" 는 OOXML의 rFonts 중
        w:ascii / w:hAnsi(영문·숫자용)만 설정하고, 한글이 속하는
        w:eastAsia는 건드리지 않음. 제목(_add_section_heading 등)에서만
        font.name을 지정하고 본문 텍스트는 지정하지 않았던 게 아니라,
        본문도 지정은 했지만(스타일 레벨) eastAsia가 비어 있다 보니
        Word/뷰어가 자체 기본 동아시아 폰트로 대체함 — 그 기본값이
        한글을 지원하지 않는 환경(웹 뷰어, 한글 폰트 없는 PC 등)에서는
        글자가 전부 네모로 보임(텍스트 데이터 자체는 깨지지 않았음).
        w:eastAsia를 명시적으로 지정해 이 문제를 구조적으로 차단함.
        """
        run.font.name = name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), name)
        r_fonts.set(qn("w:ascii"), name)
        r_fonts.set(qn("w:hAnsi"), name)

    doc = Document()

    # 기본 폰트 설정 (한글 가독성을 위해 맑은 고딕 우선)
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # style.font.name만으로는 rFonts의 w:ascii/w:hAnsi(영문용)만 설정되고
    # 한글이 속하는 w:eastAsia는 비어있는 채로 남음. 이 값이 비어있으면
    # Word/뷰어가 자체 기본 동아시아 폰트로 대체하는데, 그 기본값에 한글이
    # 없는 환경(한글 폰트 없는 웹 뷰어 등)에서는 본문 전체가 네모(□)로
    # 보이는 원인이 됨. Normal 스타일 자체에도 eastAsia를 명시해 둠
    # (아래에서 실제 각 run에도 한 번 더 명시적으로 지정해 이중으로 보강함).
    style_r_pr = style.element.get_or_add_rPr()
    style_r_fonts = style_r_pr.find(qn("w:rFonts"))
    if style_r_fonts is None:
        style_r_fonts = OxmlElement("w:rFonts")
        style_r_pr.append(style_r_fonts)
    style_r_fonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 표지 영역 ---------------------------------------------------------
    # 제목: 포인트 컬러 굵게 + 아래 강조선으로 화면 헤더의 느낌을 재현
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = PF_TEXT_STRONG
    _apply_korean_font(title_run)
    _set_paragraph_border_bottom(title_p, color_hex="2563EB", size=16)
    title_p.paragraph_format.space_after = Pt(10)

    # 메타 정보 (생성일시) - 보조 텍스트 색
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    meta_run = meta.add_run(f"생성일시  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = PF_TEXT_MUTED
    _apply_korean_font(meta_run)

    def _add_section_heading(text, level=2):
        """포인트 컬러를 쓰는 섹션 제목 (## , ### 공통 처리)"""
        h = doc.add_heading(level=level)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(text)
        run.font.color.rgb = PF_ACCENT if level == 2 else PF_TEXT_STRONG
        _apply_korean_font(run)
        run.font.size = Pt(14) if level == 2 else Pt(12)
        run.font.bold = True
        return h

    if question:
        _add_section_heading("질의 내용", level=2)
        q_p = doc.add_paragraph(question)
        q_p.paragraph_format.space_after = Pt(10)
        for r in q_p.runs:
            _apply_korean_font(r)

    _add_section_heading("AI 회신", level=2)

    for kind, content in _parse_markdown_lines(response_md):
        if kind == "blank":
            continue
        elif kind == "h3":
            _add_section_heading(content, level=3)
        elif kind == "h2":
            _add_section_heading(content, level=2)
        elif kind == "h1":
            _add_section_heading(content, level=2)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            for text, is_bold in _split_bold(content):
                run = p.add_run(text)
                run.bold = is_bold
                _apply_korean_font(run)
                if is_bold:
                    run.font.color.rgb = PF_TEXT_STRONG
        else:  # text
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            for text, is_bold in _split_bold(content):
                run = p.add_run(text)
                run.bold = is_bold
                _apply_korean_font(run)
                if is_bold:
                    run.font.color.rgb = PF_TEXT_STRONG

    # 푸터 -----------------------------------------------------------
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(20)
    _set_paragraph_border_bottom(footer_p, color_hex=PF_BORDER_GRAY, size=4)
    footer_p.paragraph_format.space_after = Pt(4)

    footer = doc.add_paragraph()
    footer_run = footer.add_run("포스원 회계법인 세무질의 AI 시스템 자동 생성")
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = PF_TEXT_MUTED
    _apply_korean_font(footer_run)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ----------------------------------------------------------------------
# PDF 변환
# ----------------------------------------------------------------------
def _find_korean_font() -> str:
    """
    한글 폰트를 탐색.

    우선순위:
    ① 이 파일(doc_export.py)과 함께 저장소에 커밋해 둔 번들 폰트
       — assets/fonts/NanumGothic-Regular.ttf 에 있어도, doc_export.py와
       같은 폴더(저장소 루트)에 바로 있어도 둘 다 찾음. 로컬/웹 배포
       환경과 무관하게 항상 존재하므로 가장 먼저 확인함. 웹 배포 환경
       (Streamlit Cloud 기본 이미지)에는 애초에 한글 폰트가 설치되어
       있지 않아 PDF 생성이 항상 실패했는데, 폰트 파일 자체를 저장소에
       포함시켜 이 문제를 근본적으로 해결함(서버에 폰트를 설치할 권한이
       없어도 됨 — 그냥 git에 파일로 들어있는 것만으로 충분).
    ② 로컬 PC(Windows)의 시스템 폰트(맑은 고딕) — 번들 폰트가 없는
       구버전 체크아웃 등 예외 상황 대비.
    ③ Linux 시스템에 나눔고딕이 별도로 설치되어 있는 경우.

    주의: .ttc(TrueType Collection, 여러 폰트가 한 파일에 묶인 형식)는 fpdf2가
    제대로 처리하지 못해 렌더링 에러가 나는 경우가 있어 후보에서 제외함.
    .ttf(단일 폰트 파일) 형식만 사용함.

    찾지 못하면 빈 문자열 반환 (이 경우 PDF 생성은 건너뛰고 md/docx로 안내).
    """
    module_dir = Path(__file__).resolve().parent
    candidates = [
        str(module_dir / "assets" / "fonts" / "NanumGothic-Regular.ttf"),  # ①-1 assets/fonts/ 폴더
        str(module_dir / "NanumGothic-Regular.ttf"),                       # ①-2 저장소 루트(doc_export.py와 같은 위치)
        r"C:\Windows\Fonts\malgun.ttf",       # ② 맑은 고딕 (로컬 PC)
        r"C:\Windows\Fonts\malgunbd.ttf",     # 맑은 고딕 Bold
        # ③ Linux (Streamlit Cloud 등 웹 서버 — 나눔고딕이 시스템에 별도 설치된 경우)
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    return ""


def build_pdf_bytes(title: str, question: str, response_md: str) -> bytes:
    """
    질의/회신을 PDF 형식의 바이트로 변환 (디스크에 쓰지 않음).
    한글 출력을 위해 시스템 한글 폰트를 사용함(없으면 RuntimeError).

    Parameters
    ----------
    title, question, response_md : build_docx_bytes와 동일

    Returns
    -------
    bytes
        PDF 파일의 바이트 내용

    Raises
    ------
    RuntimeError
        시스템에서 한글 폰트를 찾지 못한 경우 (이 환경에서는 PDF 생성 불가).
        웹 배포 환경(Streamlit Cloud 기본 이미지)에는 한글 폰트가 없는 경우가
        많으므로, 이 경우 호출하는 쪽에서 PDF를 건너뛰고 md/docx로 안내해야 함.
    """
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    font_path = _find_korean_font()
    if not font_path:
        raise RuntimeError(
            "한글 폰트를 찾을 수 없어 PDF를 생성할 수 없습니다.\n"
            "가장 확실한 해결책: NanumGothic-Regular.ttf 파일을 이 저장소의 "
            "doc_export.py와 같은 폴더(루트) 또는 assets/fonts/ 폴더 안에 "
            "커밋해두면, 로컬/웹 환경 어디서든 이 폰트가 자동으로 사용됩니다.\n"
            "(참고) 로컬 PC: Windows 폰트 폴더(C:\\Windows\\Fonts)에 malgun.ttf가 있는지 확인하세요.\n"
            "위 방법이 모두 안 되면 대신 Word(.docx) 또는 마크다운(.md) 형식으로 받아주세요."
        )

    # 디자인 테마 (streamlit_ui.py / build_docx_bytes와 색감 통일)
    PF_ACCENT = (37, 99, 235)        # #2563EB
    PF_TEXT_STRONG = (30, 58, 95)    # #1E3A5F
    PF_TEXT_MUTED = (107, 114, 128)  # #6B7280
    PF_TEXT_BODY = (51, 51, 51)      # #333333
    PF_BORDER_GRAY = (209, 213, 219)  # #D1D5DB

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("korean", "", font_path)
    pdf.set_font("korean", size=18)

    def mc(height, text):
        """
        pdf.multi_cell을 감싸는 헬퍼.

        원인 메모 (2026-07-02, 한글 폰트를 정상적으로 찾은 뒤에도 PDF 생성이
        "글자 하나 놓을 공간도 없다(Not enough horizontal space to render a
        single character)"는 오류로 실패하던 문제 조사):
        이 프로젝트가 쓰는 fpdf2 버전은 multi_cell 호출이 끝난 뒤 커서
        x좌표를 왼쪽 여백이 아니라 "방금 그린 줄의 오른쪽 끝"에 남겨두는
        것이 기본 동작(new_x 기본값 = XPos.RIGHT)임. 그 상태에서 너비를
        0(= "현재 x좌표부터 오른쪽 여백까지 남은 폭 전부 사용")으로 지정한
        다음 multi_cell을 또 호출하면, 남은 폭이 사실상 0에 가까워 글자
        하나 그릴 공간도 없다고 판단해 예외가 발생함. 이전까지는 서버에
        한글 폰트 자체가 없어 PDF 생성이 그보다 먼저 실패했기 때문에 이
        문제가 겉으로 드러나지 않았을 뿐, 폰트를 고치자 함께 드러난
        별개의 실제 버그였음.
        모든 호출을 이 헬퍼로 통일해 매번 명시적으로 왼쪽 여백으로
        커서를 되돌리도록(new_x=XPos.LMARGIN, new_y=YPos.NEXT) 고정함.
        """
        pdf.multi_cell(0, height, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # 표지 영역: 포인트 컬러 제목 + 아래 강조선 -------------------------
    pdf.set_text_color(*PF_TEXT_STRONG)
    mc(11, title)
    pdf.set_draw_color(*PF_ACCENT)
    pdf.set_line_width(0.8)
    line_y = pdf.get_y() + 1
    pdf.line(pdf.l_margin, line_y, pdf.w - pdf.r_margin, line_y)
    pdf.ln(5)

    pdf.set_font("korean", size=9)
    pdf.set_text_color(*PF_TEXT_MUTED)
    mc(6, f"생성일시  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    pdf.ln(4)

    if question:
        pdf.set_font("korean", size=13)
        pdf.set_text_color(*PF_ACCENT)
        mc(8, "질의 내용")
        pdf.set_font("korean", size=10)
        pdf.set_text_color(*PF_TEXT_BODY)
        mc(6, question)
        pdf.ln(4)

    pdf.set_font("korean", size=13)
    pdf.set_text_color(*PF_ACCENT)
    mc(8, "AI 회신")
    pdf.set_text_color(*PF_TEXT_BODY)
    pdf.ln(1)

    for kind, content in _parse_markdown_lines(response_md):
        if kind == "blank":
            pdf.ln(2)
            continue
        plain_text = content.replace("**", "")  # PDF에서는 굵게 폰트 별도 미적용(가독성 우선 단순화)
        if kind in ("h1", "h2", "h3"):
            size = 13 if kind == "h2" else 12
            pdf.set_font("korean", size=size)
            pdf.set_text_color(*PF_TEXT_STRONG)
            pdf.ln(2)
            mc(7, plain_text)
            pdf.set_font("korean", size=10)
            pdf.set_text_color(*PF_TEXT_BODY)
        elif kind == "bullet":
            mc(6, f"  - {plain_text}")
        else:
            mc(6, plain_text)

    # 푸터: 옅은 회색 구분선 + 보조 텍스트 -------------------------------
    pdf.ln(6)
    pdf.set_draw_color(*PF_BORDER_GRAY)
    pdf.set_line_width(0.3)
    footer_line_y = pdf.get_y()
    pdf.line(pdf.l_margin, footer_line_y, pdf.w - pdf.r_margin, footer_line_y)
    pdf.ln(3)
    pdf.set_font("korean", size=8)
    pdf.set_text_color(*PF_TEXT_MUTED)
    mc(5, "포스원 회계법인 세무질의 AI 시스템 자동 생성")

    return bytes(pdf.output())


# ----------------------------------------------------------------------
# 통합 함수 (메모리 바이트 반환 — 웹 다운로드 버튼 / 로컬 파일 저장 양쪽에서 재사용)
# ----------------------------------------------------------------------
def build_export_bytes(
    title: str,
    question: str,
    response_md: str,
    formats: list = None,
) -> dict:
    """
    여러 형식의 파일을 메모리(바이트)로 한 번에 생성.

    설계 의도 (2026-06-25 — 웹 배포 지원):
    - 기존 export_response()는 디스크에 직접 저장하는 방식만 지원했음. 이 함수는
      디스크에 쓰지 않고 바이트만 만들어 반환하므로, 호출하는 쪽에서
      st.download_button(웹에서 사용자 컴퓨터로 다운로드)이나, 원한다면 직접
      파일로 저장(로컬 PC)하는 데 그대로 재사용할 수 있음.
    - md는 항상 만들 수 있음. docx도 항상 만들 수 있음(폰트 의존성 없음).
      pdf는 시스템에 한글 폰트가 있어야만 만들 수 있으므로, 실패 시
      results["errors"]에 사유가 남고 results["pdf"]는 None으로 유지됨
      (다른 형식 생성에는 영향 없음).

    Returns
    -------
    dict
        {"md": bytes 또는 None, "docx": bytes 또는 None, "pdf": bytes 또는 None,
         "errors": [str, ...]}
    """
    formats = formats or ["md"]
    results = {"md": None, "docx": None, "pdf": None, "errors": []}

    if "md" in formats:
        try:
            content = (
                f"# {title}\n\n**생성일시**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
                + (f"## 질의 내용\n{question}\n\n---\n\n" if question else "")
                + f"## AI 회신\n\n{response_md}\n\n---\n*포스원 회계법인 세무질의 AI 시스템 자동 생성*\n"
            )
            results["md"] = content.encode("utf-8")
        except Exception as e:
            results["errors"].append(f"md 생성 실패: {e}")

    if "docx" in formats:
        try:
            results["docx"] = build_docx_bytes(title, question, response_md)
        except Exception as e:
            results["errors"].append(f"docx 생성 실패: {e}")

    if "pdf" in formats:
        try:
            results["pdf"] = build_pdf_bytes(title, question, response_md)
        except Exception as e:
            results["errors"].append(f"pdf 생성 실패: {e}")

    return results


def safe_filename(filename: str = None) -> str:
    """파일명에 쓸 수 없는 문자를 제거하고, 비어 있으면 타임스탬프 기반 기본값을 만듦."""
    if not filename:
        filename = f"자문_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    return re.sub(r'[\\/:*?"<>|]', "_", filename).strip() or "자문"


# ----------------------------------------------------------------------
# (하위 호환용) 로컬 PC 디스크 저장 함수
# ----------------------------------------------------------------------
def export_response(
    title: str,
    question: str,
    response_md: str,
    output_dir: Path,
    filename: str = None,
    formats: list = None,
) -> dict:
    """
    여러 형식으로 한 번에 '디스크에' 저장하는 함수. PC 로컬 환경에서만 사용 권장.

    주의: 웹 배포 환경(Streamlit Cloud 등)에서는 output_dir이 실제로 쓸 수 있는
    경로가 아닐 가능성이 높음(사용자가 PC 기준으로 적어둔 경로가 서버에는 없음).
    웹 환경에서는 이 함수 대신 build_export_bytes()로 바이트를 만들어
    st.download_button으로 사용자에게 직접 내려주는 방식을 사용해야 함.

    Parameters
    ----------
    title : str
        문서 제목
    question : str
        질문 내용 (종합 문서의 경우 빈 문자열 가능)
    response_md : str
        마크다운 회신 본문
    output_dir : Path
        저장 폴더
    filename : str, optional
        파일명(확장자 제외). None이면 "자문_YYYYMMDD_HHMMSS" 자동 생성.
    formats : list[str], optional
        ["md", "docx", "pdf"] 중 선택. None이면 ["md"]만 저장.

    Returns
    -------
    dict
        {"md": Path 또는 None, "docx": ..., "pdf": ..., "errors": [str, ...]}
    """
    filename = safe_filename(filename)
    built = build_export_bytes(title, question, response_md, formats=formats)

    results = {"md": None, "docx": None, "pdf": None, "errors": list(built["errors"])}

    try:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        # 폴더 자체를 못 만들면(웹 서버의 잘못된 경로 등) 모든 디스크 저장이
        # 의미 없으므로, 명확한 에러만 남기고 즉시 반환 (예외를 던지지 않음 —
        # 호출하는 쪽 화면이 깨지지 않도록 함)
        results["errors"].append(f"저장 폴더를 만들 수 없습니다 ({output_dir}): {e}")
        return results

    ext_map = {"md": ".md", "docx": ".docx", "pdf": ".pdf"}
    for fmt, ext in ext_map.items():
        data = built.get(fmt)
        if data is None:
            continue
        try:
            file_path = output_dir / f"{filename}{ext}"
            file_path.write_bytes(data)
            results[fmt] = file_path
        except Exception as e:
            results["errors"].append(f"{fmt} 저장 실패: {e}")

    return results
